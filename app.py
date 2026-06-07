"""
JobMate (Cowork edition) — Flask API backed by the local Claude Agent SDK
Run: python app.py
Requires: pip install -r requirements.txt
Auth: uses your local Claude Code login — no ANTHROPIC_API_KEY needed.
"""

import asyncio
import io
import json
import os
import re
import threading

from flask import Flask, jsonify, request, send_from_directory, send_file
from flask_cors import CORS

from claude_agent_sdk import ClaudeSDKClient, ClaudeAgentOptions, query
from claude_agent_sdk.types import AssistantMessage, TextBlock

app = Flask(__name__, static_folder="static")
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB upload cap


@app.errorhandler(Exception)
def _json_error(e):
    """Return JSON for any unhandled exception so the frontend can show it."""
    import traceback
    traceback.print_exc()  # still log full trace to terminal
    code = getattr(e, "code", 500) if isinstance(getattr(e, "code", None), int) else 500
    return jsonify({"error": f"{type(e).__name__}: {e}"}), code

SYSTEM_PROMPT = (
    "You are a senior career coach and ATS specialist. "
    "When the user asks for JSON, return ONLY a single valid JSON object — "
    "no prose, no markdown fences, no commentary."
)

# ── Persistent Claude SDK client ─────────────────────────────────────────────
# One long-lived ClaudeSDKClient lives on a background event loop, shared
# across all Flask requests. This avoids spawning a fresh `claude` subprocess
# (and re-paying ~11k tokens of cache creation) on every call.
#
# Trade-off: ClaudeSDKClient preserves conversation history across queries.
# Each JobMate prompt is self-contained (it includes the full resume + JD), so
# context bleed is minimal — but input tokens do grow over the app's lifetime.
# Restart Flask if you want a clean slate.

_loop = None
_client = None
_init_lock = threading.Lock()      # guard one-time loop/client setup
_request_lock = threading.Lock()   # serialize concurrent Flask requests against the client


async def _setup_client():
    global _client
    options = ClaudeAgentOptions(
        allowed_tools=[],
        system_prompt=SYSTEM_PROMPT,
    )
    _client = ClaudeSDKClient(options=options)
    await _client.connect()


def _ensure_client():
    """Lazily start the background event loop and connect the SDK client."""
    global _loop
    if _client is not None:
        return
    with _init_lock:
        if _client is not None:
            return
        ready = threading.Event()
        err_box = {}

        def _runner():
            global _loop
            _loop = asyncio.new_event_loop()
            asyncio.set_event_loop(_loop)
            try:
                _loop.run_until_complete(_setup_client())
            except Exception as e:
                err_box["error"] = e
            finally:
                ready.set()
            _loop.run_forever()

        threading.Thread(target=_runner, daemon=True, name="claude-sdk-loop").start()
        if not ready.wait(timeout=60):
            raise RuntimeError("Claude SDK client failed to initialize within 60s")
        if "error" in err_box:
            raise err_box["error"]


async def _ask(prompt: str) -> str:
    """Send one prompt on the persistent client and return concatenated text."""
    await _client.query(prompt)
    out = []
    async for message in _client.receive_response():
        if isinstance(message, AssistantMessage):
            for block in message.content:
                if isinstance(block, TextBlock):
                    out.append(block.text)
    return "".join(out)


def claude(prompt: str, max_tokens: int = 2500) -> dict:
    """Call Claude via the persistent client and parse the JSON response.

    `max_tokens` kept for signature compatibility with the prior API build; the
    SDK manages its own output limits.
    """
    _ensure_client()
    # The client is single-threaded — one in-flight query at a time.
    with _request_lock:
        future = asyncio.run_coroutine_threadsafe(_ask(prompt), _loop)
        text = future.result(timeout=300)

    match = re.search(r"```json\s*(.*?)```", text, re.DOTALL) or re.search(r"(\{.*\})", text, re.DOTALL)
    if match:
        return json.loads(match.group(1))
    return json.loads(text)


# ── File extractors ──────────────────────────────────────────────────────────

def extract_pdf(stream: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(stream))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def extract_docx(stream: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(stream))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


# ── Resume upload ────────────────────────────────────────────────────────────

@app.route("/api/resume/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    name = f.filename.lower()
    data = f.read()

    try:
        if name.endswith(".pdf"):
            text = extract_pdf(data)
        elif name.endswith(".docx"):
            text = extract_docx(data)
        elif name.endswith((".txt", ".md")):
            text = data.decode("utf-8", errors="ignore").strip()
        else:
            return jsonify({"error": "Unsupported format. Use PDF, DOCX, TXT, or MD."}), 415
    except Exception as e:
        return jsonify({"error": f"Could not parse {f.filename}: {e}"}), 422

    if not text or len(text) < 50:
        return jsonify({"error": "Extracted text is too short — file may be image-only or empty."}), 422

    return jsonify({
        "filename": f.filename,
        "text": text,
        "wordCount": len(text.split()),
        "charCount": len(text),
    })


# ── Step 1: Score & gap analysis ─────────────────────────────────────────────

@app.route("/api/analyze", methods=["POST"])
def analyze():
    d = request.json
    resume, jd = d.get("resume", ""), d.get("jobDescription", "")
    if not resume or not jd:
        return jsonify({"error": "Resume and job description are required"}), 400

    result = claude(f"""
You are a senior recruiter and ATS specialist. Analyze this resume against the job description.

RESUME:
{resume}

JOB DESCRIPTION:
{jd}

Return ONLY valid JSON (no prose) matching this schema exactly:
{{
  "scores": {{
    "keywordOverlap": <int 0-100>,
    "skillsMatch": <int 0-100>,
    "outcomes": <int 0-100>,
    "roleFit": <int 0-100>,
    "overall": <int 0-100>
  }},
  "missingKeywords": ["string"],
  "weakBullets": [
    {{"original": "string", "issue": "string", "suggestion": "string"}}
  ],
  "atsIssues": [
    {{"section": "string", "issue": "string", "fix": "string"}}
  ],
  "gaps": ["string"],
  "strengths": ["string"],
  "summary": "string",
  "wouldInterview": true,
  "interviewReason": "string"
}}
""")
    return jsonify(result)


# ── Step 2: Rewrite bullets (XYZ formula) ────────────────────────────────────

@app.route("/api/tailor/bullets", methods=["POST"])
def tailor_bullets():
    d = request.json
    resume, jd = d.get("resume", ""), d.get("jobDescription", "")

    result = claude(f"""
You are a senior career coach. Rewrite every bullet point using Google's XYZ formula:
"Accomplished X as measured by Y by doing Z."
Rules: under 20 words, metric-driven, use the JD's exact language, no invented numbers.
Flag estimates with "(approx.)".

RESUME: {resume}
JOB DESCRIPTION: {jd}

Return ONLY valid JSON:
{{
  "rewrites": [
    {{"original": "string", "rewritten": "string", "formula": "X / Y / Z breakdown"}}
  ]
}}
""")
    return jsonify(result)


# ── Step 3: ATS audit & fixes ─────────────────────────────────────────────────

@app.route("/api/tailor/ats", methods=["POST"])
def tailor_ats():
    d = request.json
    resume = d.get("resume", "")

    result = claude(f"""
You are an ATS expert. Audit this resume for parsing failures.
Check: special characters (pipes, em dashes), column layouts, header/footer content,
skills section formatting, non-standard section names, foreign degree recognition,
keyword density gaps.

RESUME: {resume}

Return ONLY valid JSON:
{{
  "issues": [
    {{
      "severity": "critical|warning",
      "section": "string",
      "problem": "string",
      "fix": "string",
      "before": "string",
      "after": "string"
    }}
  ],
  "overallAtsScore": <int 0-100>,
  "summary": "string"
}}
""")
    return jsonify(result)


# ── Step 4: Hiring manager review ─────────────────────────────────────────────

@app.route("/api/tailor/hm-review", methods=["POST"])
def hm_review():
    d = request.json
    resume, jd = d.get("resume", ""), d.get("jobDescription", "")

    result = claude(f"""
Act as a hiring manager skimming this resume for 10 seconds against the job description.
Score on keyword overlap, skills, outcomes, and role fit.
Identify weak bullets, missing high-priority terms, and exactly what to change to reach 80%.
State honestly whether you would interview this candidate.

RESUME: {resume}
JOB DESCRIPTION: {jd}

Return ONLY valid JSON:
{{
  "scores": {{
    "keywordOverlap": <int>,
    "skillsMatch": <int>,
    "outcomes": <int>,
    "roleFit": <int>,
    "overall": <int>
  }},
  "missingTerms": [{{"term": "string", "priority": "high|medium", "fix": "string"}}],
  "weakBullets": [{{"bullet": "string", "problem": "string", "rewrite": "string"}}],
  "changesToClear80": ["string"],
  "wouldInterview": true,
  "interviewReason": "string",
  "topStrength": "string",
  "topConcern": "string"
}}
""")
    return jsonify(result)


# ── Step 4b: Produce final tailored resume ────────────────────────────────────

@app.route("/api/tailor/finalize", methods=["POST"])
def finalize():
    d = request.json
    resume, jd = d.get("resume", ""), d.get("jobDescription", "")

    result = claude(f"""
You are a senior career coach. Produce the complete optimized resume, incorporating:
- XYZ formula bullets
- ATS-safe formatting (no pipes, no em dashes, plain ASCII bullets)
- JD keyword alignment
- NIST/framework alignment where applicable
- All fixes from prior audit steps

RESUME: {resume}
JOB DESCRIPTION: {jd}

Return ONLY valid JSON:
{{
  "tailoredResume": "full resume as plain text",
  "changesSummary": ["string"],
  "finalScore": <int 0-100>
}}
""")
    return jsonify(result)


# ── Job search ────────────────────────────────────────────────────────────────

@app.route("/api/jobs/search", methods=["POST"])
def search_jobs():
    d = request.json
    resume = d.get("resume", "")
    location = d.get("location", "Calgary, AB")
    remote = d.get("includeRemote", True)

    result = claude(f"""
Based on this resume, identify the 20 best-fit job roles to target.
Consider experience level, skills, certifications, and location ({location}).
{"Include remote-friendly roles across Canada." if remote else ""}

RESUME: {resume}

Return ONLY valid JSON:
{{
  "jobs": [
    {{
      "id": "unique_string",
      "title": "string",
      "company": "string",
      "location": "string",
      "matchScore": <int 0-100>,
      "tier": <1|2|3>,
      "tierLabel": "Best Shot|Strong Match|Stretch",
      "reason": "string",
      "missingSkills": ["string"],
      "salary": "string",
      "linkedinUrl": "https://ca.linkedin.com/jobs/search/?keywords=...",
      "indeedUrl": "https://ca.indeed.com/q-...",
      "requiresCoverLetter": true
    }}
  ]
}}
""")
    return jsonify(result)


# ── Cover letter ──────────────────────────────────────────────────────────────

@app.route("/api/cover-letter", methods=["POST"])
def cover_letter():
    d = request.json
    resume = d.get("resume", "")
    jd = d.get("jobDescription", "")
    company = d.get("company", "the company")
    role = d.get("role", "this role")
    user_name = d.get("userName", "Candidate")

    result = claude(f"""
Write a compelling, personalized cover letter for {user_name} applying to {role} at {company}.

Rules:
- 3 paragraphs, under 350 words
- Lead with the strongest metric from the resume
- Name at least 2 specific tools/frameworks from the JD
- Second paragraph: connect 2-3 experience highlights to JD requirements
- Third paragraph: mention Calgary base, availability, enthusiasm
- Professional but not robotic

RESUME: {resume}
JOB DESCRIPTION: {jd}

Return ONLY valid JSON:
{{
  "coverLetter": "full letter text with \\n line breaks",
  "wordCount": <int>,
  "keyToolsNamed": ["string"]
}}
""")
    return jsonify(result)


# ── Application form helper ───────────────────────────────────────────────────

@app.route("/api/apply/extract-fields", methods=["POST"])
def extract_fields():
    d = request.json
    resume = d.get("resume", "")
    form_fields = d.get("formFields", [])

    result = claude(f"""
Given this resume and a list of job application form fields, determine what can be
auto-filled and what the user must answer manually.

RESUME: {resume}

FORM FIELDS: {json.dumps(form_fields)}

Return ONLY valid JSON:
{{
  "autoFilled": [
    {{"field": "string", "value": "string", "confidence": "high|medium|low"}}
  ],
  "needsUserInput": [
    {{"field": "string", "question": "How should I phrase this for you?", "type": "text|select|yesno", "options": []}}
  ]
}}
""")
    return jsonify(result)


# ── Fetch JD from a posting URL ──────────────────────────────────────────────
# Uses an ephemeral query() (not the persistent client) so we can enable
# WebFetch for this one call without changing the long-lived client's options.

@app.route("/api/jobs/fetch-description", methods=["POST"])
def fetch_job_description():
    """Fetch a job description either from a known URL or by searching for the
    posting given {title, company, location}. The latter is what the Apply flow
    uses so the user doesn't have to paste anything."""
    d = request.json or {}
    url = (d.get("url") or "").strip()
    job = d.get("job") or {}
    title = (job.get("title") or "").strip()
    company = (job.get("company") or "").strip()
    location = (job.get("location") or "").strip()

    if not url and not (title and company):
        return jsonify({"error": "Provide either `url` or `job: {title, company, location}`"}), 400

    if url:
        if not (url.startswith("http://") or url.startswith("https://")):
            return jsonify({"error": "url must start with http:// or https://"}), 400
        tools = ["WebFetch"]
        instruction = f"Fetch this job posting URL and extract the details:\n{url}"
    else:
        tools = ["WebSearch", "WebFetch"]
        loc_clause = f" in {location}" if location else ""
        instruction = (
            f"Find the current job posting for \"{title}\" at {company}{loc_clause}.\n"
            f"1. Use WebSearch to locate it. Try the company's own careers page first, "
            f"then job boards (LinkedIn, Indeed, Glassdoor, Lever, Greenhouse).\n"
            f"2. Use WebFetch on the best match to extract the full job description.\n"
            f"3. If you can't find an exact title match, return the closest current "
            f"opening at the same company that fits the role.\n"
            f"4. If no current posting exists, set fetchOk=false and explain in note."
        )

    async def _fetch():
        opts = ClaudeAgentOptions(
            allowed_tools=tools,
            system_prompt=(
                "You extract job posting details. "
                "Return ONLY a JSON object — no prose, no markdown fences."
            ),
        )
        prompt = f"""{instruction}

If the page requires login or is blocked, return the best information you can
infer from any accessible content. Set "fetchOk" to false if you could not get
the actual posting.

Return ONLY valid JSON:
{{
  "company": "string",
  "role": "string",
  "location": "string",
  "postingUrl": "the URL of the actual posting you found and fetched (empty string if none)",
  "jobDescription": "the full job description as plain text",
  "fetchOk": true,
  "note": "short note if anything is approximate, missing, or if the source was a cached/secondary listing"
}}"""
        out = []
        async for message in query(prompt=prompt, options=opts):
            if isinstance(message, AssistantMessage):
                for block in message.content:
                    if isinstance(block, TextBlock):
                        out.append(block.text)
        return "".join(out)

    text = asyncio.run(_fetch())
    match = re.search(r"```json\s*(.*?)```", text, re.DOTALL) or re.search(r"(\{.*\})", text, re.DOTALL)
    if match:
        result = json.loads(match.group(1))
    else:
        result = json.loads(text)
    return jsonify(result)


# ── Resume export (PDF & DOCX) ───────────────────────────────────────────────

def _classify_line(stripped: str):
    """Return ('header'|'bullet'|'body', cleaned_text)."""
    if not stripped:
        return ("blank", "")
    if stripped.startswith("#"):
        return ("header", stripped.lstrip("#").strip())
    # ALL-CAPS short line → section header
    letters = [c for c in stripped if c.isalpha()]
    if letters and all(c.isupper() for c in letters) and 2 < len(stripped) <= 60:
        return ("header", stripped)
    if stripped[0] in "-*•·":
        return ("bullet", stripped.lstrip("-*•· ").strip())
    return ("body", stripped)


def text_to_pdf(text: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF(format="Letter", unit="pt")
    pdf.set_auto_page_break(True, margin=36)
    pdf.add_page()
    pdf.set_margins(48, 36, 48)
    pdf.set_font("Helvetica", "", 10.5)

    for raw in text.split("\n"):
        line = raw.rstrip()
        # fpdf2 default fonts (Helvetica) are latin-1; replace anything that isn't.
        safe = line.encode("latin-1", "replace").decode("latin-1")
        kind, clean = _classify_line(safe.strip())

        if kind == "blank":
            pdf.ln(6)
        elif kind == "header":
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 14, clean)
            pdf.set_font("Helvetica", "", 10.5)
        elif kind == "bullet":
            pdf.multi_cell(0, 12, "- " + clean)
        else:
            pdf.multi_cell(0, 12, clean)

    return bytes(pdf.output())


def text_to_docx(text: str) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for raw in text.split("\n"):
        line = raw.rstrip()
        kind, clean = _classify_line(line.strip())

        if kind == "blank":
            doc.add_paragraph()
        elif kind == "header":
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(12)
        elif kind == "bullet":
            doc.add_paragraph(clean, style="List Bullet")
        else:
            doc.add_paragraph(clean)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/api/export/pdf", methods=["POST"])
def export_pdf():
    d = request.json or {}
    text = d.get("text", "")
    filename = d.get("filename") or "tailored_resume.pdf"
    if not text.strip():
        return jsonify({"error": "text is required"}), 400
    return send_file(
        io.BytesIO(text_to_pdf(text)),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    d = request.json or {}
    text = d.get("text", "")
    filename = d.get("filename") or "tailored_resume.docx"
    if not text.strip():
        return jsonify({"error": "text is required"}), 400
    return send_file(
        io.BytesIO(text_to_docx(text)),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ── Serve frontend ────────────────────────────────────────────────────────────

@app.route("/", defaults={"path": ""})
@app.route("/<path:path>")
def serve(path):
    if path and os.path.exists(os.path.join(app.static_folder, path)):
        return send_from_directory(app.static_folder, path)
    return send_from_directory(app.static_folder, "index.html")


if __name__ == "__main__":
    print("JobMate (Cowork) running → http://localhost:5000")
    print("Backed by your local Claude Code login. No API key required.")
    app.run(debug=True, port=5000)
